"""申诉函件 Word 生成器"""
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class AppealLetterGenerator:
    """Word申诉函件生成器"""

    def __init__(self):
        self.attribution_labels = {
            '规格型号差': '规格型号不一致',
            '数量差': '数量不一致',
            '单价差': '单价不一致',
            '项目对照差': '项目对照不一致',
            '其他': '其他差异'
        }

    def generate(self, diff_records: List[Dict[str, Any]], output_path: str) -> None:
        """生成申诉函件草稿

        Args:
            diff_records: 差异记录列表
            output_path: 输出文件路径
        """
        doc = Document()

        # 标题
        title = doc.add_heading('医保结算差异申诉函', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"日期: {datetime.now().strftime('%Y年%m月%d日')}")

        # 差异汇总统计
        stats = self._collect_statistics(diff_records)

        doc.add_paragraph()
        summary = doc.add_paragraph()
        summary.add_run('一、差异汇总').bold = True

        doc.add_paragraph(f"本次对账共发现 {len(diff_records)} 条差异记录，分类汇总如下：")

        # 统计表
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '差异类型'
        hdr_cells[1].text = '数量'
        hdr_cells[2].text = '说明'

        for attr_type, count in stats['by_type'].items():
            row_cells = table.add_row().cells
            label = self.attribution_labels.get(attr_type, attr_type)
            row_cells[0].text = label
            row_cells[1].text = str(count)
            row_cells[2].text = ''

        doc.add_paragraph()

        # 差异明细
        detail = doc.add_paragraph()
        detail.add_run('二、差异明细').bold = True

        for idx, record in enumerate(diff_records[:20], 1):  # 最多20条
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run(f"{idx}. ").bold = True
            p.add_run(f"[{record.get('attribution_type', '未知')}] ")
            p.add_run(record.get('attribution_reason', ''))

            diff_values = record.get('diff_values', {})
            if diff_values:
                p2 = doc.add_paragraph()
                p2.add_run(f"   差异详情: {diff_values}")

        if len(diff_records) > 20:
            doc.add_paragraph(f"... (共 {len(diff_records)} 条，仅显示前20条)")

        # 签名区
        doc.add_paragraph()
        doc.add_paragraph()
        sign_area = doc.add_paragraph()
        sign_area.add_run('三、附件清单').bold = True

        doc.add_paragraph("1. 差异台账 (diffLedger.xlsx)")
        doc.add_paragraph("2. HTML差异报告 (diffReport.html)")

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("申请人签名: _______________")
        doc.add_paragraph("审核人签名: _______________")

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("（本函件为草稿，仅供参考）")

        doc.save(output_path)

    def _collect_statistics(self, diff_records: List[Dict[str, Any]]) -> Dict[str, Any]:
        """收集差异统计信息"""
        stats = {
            'total': len(diff_records),
            'by_type': {}
        }

        for record in diff_records:
            attr_type = record.get('attribution_type', '未知')
            stats['by_type'][attr_type] = stats['by_type'].get(attr_type, 0) + 1

        return stats
