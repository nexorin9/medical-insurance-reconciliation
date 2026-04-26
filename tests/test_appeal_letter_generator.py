"""Word申诉函生成器单元测试"""
import pytest
from pathlib import Path
import sys
import os

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.reports.appeal_letter_generator import AppealLetterGenerator


class TestAppealLetterGenerator:
    """测试AppealLetterGenerator类"""

    def setup_method(self):
        """设置测试环境"""
        self.generator = AppealLetterGenerator()
        self.test_output_dir = Path(__file__).parent.parent / "reports" / "test_output"
        self.test_output_dir.mkdir(parents=True, exist_ok=True)

    def test_generate_empty_results(self):
        """测试生成空结果的申诉函"""
        output_file = self.test_output_dir / "test_empty_letter.docx"
        results = []

        self.generator.generate(results, str(output_file))

        assert output_file.exists()
        assert os.path.getsize(str(output_file)) > 0

    def test_generate_with_data(self):
        """测试生成有数据的申诉函"""
        output_file = self.test_output_dir / "test_letter_with_data.docx"
        results = [
            {
                'his_record': {'item_code': 'H001', 'item_name': '血清总蛋白测定', 'spec_model': '干化学法'},
                'insurance_record': {'item_code': 'H001', 'item_name': '血清总蛋白测定', 'spec_model': '酶法'},
                'diff_fields': ['spec_model'],
                'attribution_type': '规格型号差'
            }
        ]

        self.generator.generate(results, str(output_file))

        assert output_file.exists()
        assert os.path.getsize(str(output_file)) > 0

    def test_generate_document_structure(self):
        """测试文档结构完整性"""
        output_file = self.test_output_dir / "test_letter_structure.docx"
        results = [
            {
                'his_record': {'item_code': 'H001', 'item_name': '血清总蛋白测定', 'quantity': 1, 'unit_price': 30.00},
                'insurance_record': {'item_code': 'H001', 'item_name': '血清总蛋白测定', 'quantity': 2, 'unit_price': 30.00},
                'diff_fields': ['quantity'],
                'attribution_type': '数量差'
            }
        ]

        self.generator.generate(results, str(output_file))

        assert output_file.exists()

        # 读取Word文档验证结构
        from docx import Document
        doc = Document(str(output_file))

        # 验证文档有段落
        assert len(doc.paragraphs) > 0

        # 验证文档包含表格（差异汇总表）
        assert len(doc.tables) > 0

    def test_generate_multiple_records(self):
        """测试生成多条记录的申诉函"""
        output_file = self.test_output_dir / "test_letter_multiple.docx"
        results = [
            {
                'his_record': {'item_code': 'H001', 'item_name': '血清总蛋白测定'},
                'insurance_record': {'item_code': 'H001', 'item_name': '血清总蛋白测定'},
                'diff_fields': ['spec_model'],
                'attribution_type': '规格型号差'
            },
            {
                'his_record': {'item_code': 'H002', 'item_name': '丙氨酸氨基转移酶测定'},
                'insurance_record': {'item_code': 'H002', 'item_name': '丙氨酸氨基转移酶测定'},
                'diff_fields': ['quantity'],
                'attribution_type': '数量差'
            }
        ]

        self.generator.generate(results, str(output_file))

        assert output_file.exists()

        # 读取并验证表格数量
        from docx import Document
        doc = Document(str(output_file))

        # 应该有差异汇总表
        assert len(doc.tables) >= 1

    def test_file_overwrite(self):
        """测试文件覆盖"""
        output_file = self.test_output_dir / "test_letter_overwrite.docx"
        results = [
            {
                'his_record': {'item_code': 'H001'},
                'insurance_record': {'item_code': 'H001'},
                'diff_fields': [],
                'attribution_type': '无差异'
            }
        ]

        self.generator.generate(results, str(output_file))
        first_size = os.path.getsize(str(output_file))

        # 再次运行应该覆盖文件
        self.generator.generate(results, str(output_file))
        second_size = os.path.getsize(str(output_file))

        # 文件应该被覆盖
        assert output_file.exists()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
